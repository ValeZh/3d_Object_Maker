from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def set_default_style(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(1.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(14)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(max(16 - level, 12))
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)


def add_normal_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = p.paragraph_format
    fmt.first_line_indent = Cm(1.25)
    fmt.line_spacing = 1.5
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)


def add_list_item(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = p.paragraph_format
    fmt.left_indent = Cm(0.63)
    fmt.first_line_indent = Cm(0)
    fmt.line_spacing = 1.5
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)


def add_code_line(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt = p.paragraph_format
    fmt.left_indent = Cm(1.0)
    fmt.first_line_indent = Cm(0)
    fmt.line_spacing = 1.0
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)


def convert_markdown_to_docx(md_path: Path, docx_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    set_default_style(doc)

    in_code = False
    for raw in lines:
        line = raw.rstrip()
        if line.startswith("```"):
            in_code = not in_code
            continue

        if in_code:
            add_code_line(doc, line)
            continue

        if not line.strip():
            doc.add_paragraph("")
            continue

        if line.startswith("---"):
            p = doc.add_paragraph("")
            p.paragraph_format.space_after = Pt(4)
            continue

        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            add_heading(doc, line[level:].strip(), level)
            continue

        stripped = line.lstrip()
        if stripped.startswith("- "):
            add_list_item(doc, stripped[2:].strip())
            continue

        add_normal_paragraph(doc, line)

    doc.save(docx_path)


if __name__ == "__main__":
    root = Path(__file__).resolve().parents[1]
    md = root / "src" / "generator" / "procedural" / "DIPLOMA_PROCEDURAL_DESCRIPTION.md"
    out = root / "src" / "generator" / "procedural" / "DIPLOMA_PROCEDURAL_DESCRIPTION.docx"
    convert_markdown_to_docx(md, out)
    print(f"[OK] Created: {out}")
